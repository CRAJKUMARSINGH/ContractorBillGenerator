import streamlit as st
import pandas as pd
import pdfkit
from docx import Document
import os
import zipfile
import tempfile
from jinja2 import Environment, FileSystemLoader, TemplateNotFound
from pypdf import PdfReader, PdfWriter
import numpy as np
from datetime import datetime
import traceback
import shutil
from num2words import num2words

# Initialize Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)
env.filters['strptime'] = lambda s, fmt: datetime.strptime(s, fmt) if s else None
TEMP_DIR = tempfile.mkdtemp()

# Configure wkhtmltopdf
wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)

def number_to_words(number):
    try:
        return num2words(int(number), lang="en_IN").title()
    except:
        return str(number)

def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type, amount_paid_last_bill, is_first_bill, is_final_bill, user_inputs):
    st.write("Starting process_bill")
    first_page_data = {"header": {}, "items": [], "totals": {}}
    certificate_ii_data = {"payable_amount": 0, "amount_words": "", "summary": {}}
    certificate_iii_data = {"payable_amount": 0, "amount_words": "", "summary": {}, "certification": "Certified that the work has been completed as per specifications."}
    deviation_data = {"items": [], "summary": {}} if is_final_bill else None
    extra_items_data = {"items": [], "totals": {"payable": 0}}
    note_sheet_data = {
        "notes": [],
        "header": {},
        "totals": {},
        "work_order_amount": user_inputs.get("work_order_amount", 854678.0)  # Ensure this matches user input
    }

    # Header data from user inputs
    first_page_data["header"] = {
        "agreement_no": user_inputs.get("agreement_no", "48/2024-25"),
        "name_of_work": user_inputs.get("name_of_work", "Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur"),
        "name_of_firm": user_inputs.get("name_of_firm", "M/s Seema Electrical Udaipur"),
        "date_commencement": user_inputs.get("date_commencement", "18/01/2025"),
        "date_completion": user_inputs.get("date_completion", "17/04/2025"),
        "actual_completion": user_inputs.get("actual_completion", "01/03/2025"),
        "serial_no_bill": user_inputs.get("serial_no_bill", "First & Final Bill"),
        "work_order_ref": user_inputs.get("work_order_ref", "1179 dated 09-01-2025"),
        "measurement_date": user_inputs.get("measurement_date", "03/03/2025"),
        "work_order_amount": user_inputs.get("work_order_amount", 854678.0)
    }
    note_sheet_data["header"] = first_page_data["header"].copy()

    # Log sheet shapes and sample data
    print(f"Work Order shape: {ws_wo.shape}")
    print(f"Work Order sample (rows 21-23):\n{ws_wo.iloc[20:23].to_string()}")
    print(f"Bill Quantity shape: {ws_bq.shape}")
    print(f"Bill Quantity sample (rows 21-23):\n{ws_bq.iloc[20:23].to_string()}")
    print(f"Extra Items shape: {ws_extra.shape}")
    print(f"Extra Items sample (rows 6-8):\n{ws_extra.iloc[5:8].to_string() if ws_extra.shape[0] >= 8 else 'Not enough rows or empty'}")

    # Validate sheets
    if ws_wo.empty or ws_bq.empty:
        raise ValueError("Work Order or Bill Quantity sheet is empty")
    if ws_wo.shape[0] < 22 or ws_bq.shape[0] < 22:
        raise ValueError("Work Order or Bill Quantity sheet has insufficient rows (need at least 22)")

    # Work Order items (start from row 22, 0-based index 21)
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        qty_raw = ws_bq.iloc[i, 2] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 2]) else None
        rate_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                print(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                continue

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                print(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                continue

        item = {
            "serial_no": str(i - 20),
            "description": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "unit": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "quantity": qty,
            "rate": rate,
            "amount": round(qty * rate) if qty and rate else 0,
            "bsr": str(ws_wo.iloc[i, 5]) if pd.notnull(ws_wo.iloc[i, 5]) else "",
            "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
            "is_divider": False
        }
        first_page_data["items"].append(item)

    # Extra Items processing (optional)
    if not ws_extra.empty and ws_extra.shape[0] >= 7:
        first_page_data["items"].append({
            "description": "Extra Items (With Premium)",
            "bold": True,
            "underline": True,
            "amount": 0,
            "quantity": 0,
            "rate": 0,
            "serial_no": "",
            "unit": "",
            "bsr": "",
            "remark": "",
            "is_divider": True
        })

        last_row_extra = ws_extra.shape[0]
        for j in range(6, last_row_extra):
            qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
            rate_raw = ws_extra.iloc[j, 4] if pd.notnull(ws_extra.iloc[j, 4]) else None
            amount_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

            qty = 0
            if isinstance(qty_raw, (int, float)):
                qty = float(qty_raw)
            elif isinstance(qty_raw, str):
                cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
                try:
                    qty = float(cleaned_qty)
                except ValueError:
                    print(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                    continue

            rate = 0
            if isinstance(rate_raw, (int, float)):
                rate = float(rate_raw)
            elif isinstance(rate_raw, str):
                cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
                try:
                    rate = float(cleaned_rate)
                except ValueError:
                    print(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                    continue

            amount = 0
            if isinstance(amount_raw, (int, float)):
                amount = float(amount_raw)
            elif isinstance(amount_raw, str):
                cleaned_amount = amount_raw.strip().replace(',', '').replace(' ', '')
                try:
                    amount = float(cleaned_amount)
                except ValueError:
                    print(f"Skipping invalid amount at Extra Items row {j+1}: '{amount_raw}'")
                    continue

            item = {
                "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else str(j - 5),
                "ref_bsr": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
                "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
                "unit": "",
                "quantity": qty,
                "rate": rate,
                "amount": amount if amount else round(qty * rate) if qty and rate else 0,
                "remark": str(ws_extra.iloc[j, 6]) if pd.notnull(ws_extra.iloc[j, 6]) else "",
                "is_divider": False
            }
            first_page_data["items"].append(item)
            extra_items_data["items"].append(item.copy())
            extra_items_data["totals"]["payable"] += item["amount"]

    # Totals
    data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
    total_amount = round(sum(item.get("amount", 0) for item in data_items))
    premium_amount = round(total_amount * (premium_percent / 100) if premium_type == "above" else -total_amount * (premium_percent / 100))
    original_payable = round(total_amount + premium_amount)
    payable_amount = round(original_payable - (amount_paid_last_bill if not is_first_bill else 0))

    first_page_data["totals"] = {
        "grand_total": total_amount,
        "premium": {"percent": premium_percent / 100, "type": premium_type, "amount": premium_amount},
        "original_payable": original_payable,
        "payable": payable_amount,
        "amount_paid_last_bill": amount_paid_last_bill if not is_first_bill else 0,
        "extra_items_sum": extra_items_data["totals"]["payable"],
        "extra_items_total": extra_items_data["totals"]["payable"]
    }
    note_sheet_data["totals"] = first_page_data["totals"].copy()

    # Certificate II and III
    certificate_ii_data = {
        "payable_amount": payable_amount,
        "amount_words": number_to_words(payable_amount),
        "summary": first_page_data["totals"].copy()
    }
    certificate_iii_data = {
        "payable_amount": payable_amount,
        "amount_words": number_to_words(payable_amount),
        "summary": first_page_data["totals"].copy(),
        "certification": "Certified that the work has been completed as per specifications."
    }

    # Deviation Statement (only for final bill)
    if is_final_bill:
        last_row_wo = ws_wo.shape[0]
        work_order_total = 0
        executed_total = 0
        overall_excess = 0
        overall_saving = 0
        for i in range(21, last_row_wo):
            qty_wo_raw = ws_wo.iloc[i, 2] if pd.notnull(ws_wo.iloc[i, 2]) else None
            rate_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else None
            qty_bill_raw = ws_bq.iloc[i, 2] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 2]) else None

            qty_wo = 0
            if isinstance(qty_wo_raw, (int, float)):
                qty_wo = float(qty_wo_raw)
            elif isinstance(qty_wo_raw, str):
                cleaned_qty_wo = qty_wo_raw.strip().replace(',', '').replace(' ', '')
                try:
                    qty_wo = float(cleaned_qty_wo)
                except ValueError:
                    print(f"Skipping invalid qty_wo at row {i+1}: '{qty_wo_raw}'")
                    continue

            rate = 0
            if isinstance(rate_raw, (int, float)):
                rate = float(rate_raw)
            elif isinstance(rate_raw, str):
                cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
                try:
                    rate = float(cleaned_rate)
                except ValueError:
                    print(f"Skipping invalid rate at row {i+1}: '{rate_raw}'")
                    continue

            qty_bill = 0
            if isinstance(qty_bill_raw, (int, float)):
                qty_bill = float(qty_bill_raw)
            elif isinstance(qty_bill_raw, str):
                cleaned_qty_bill = qty_bill_raw.strip().replace(',', '').replace(' ', '')
                try:
                    qty_bill = float(cleaned_qty_bill)
                except ValueError:
                    print(f"Skipping invalid qty_bill at row {i+1}: '{qty_bill_raw}'")
                    continue

            amt_wo = round(qty_wo * rate)
            amt_bill = round(qty_bill * rate)
            excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
            excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
            saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
            saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

            item = {
                "serial_no": str(i - 20),
                "description": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "unit": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "qty_wo": qty_wo,
                "rate": rate,
                "amt_wo": amt_wo,
                "qty_bill": qty_bill,
                "amt_bill": amt_bill,
                "excess_qty": excess_qty,
                "excess_amt": excess_amt,
                "saving_qty": saving_qty,
                "saving_amt": saving_amt,
                "bsr": str(ws_wo.iloc[i, 5]) if pd.notnull(ws_wo.iloc[i, 5]) else ""
            }
            deviation_data["items"].append(item)
            work_order_total += amt_wo
            executed_total += amt_bill
            overall_excess += excess_amt
            overall_saving += saving_amt

        # Deviation Summary
        tender_premium_f = round(work_order_total * (premium_percent / 100) if premium_type == "above" else -work_order_total * (premium_percent / 100))
        tender_premium_h = round(executed_total * (premium_percent / 100) if premium_type == "above" else -executed_total * (premium_percent / 100))
        tender_premium_j = round(overall_excess * (premium_percent / 100) if premium_type == "above" else -overall_excess * (premium_percent / 100))
        tender_premium_l = round(overall_saving * (premium_percent / 100) if premium_type == "above" else -overall_saving * (premium_percent / 100))
        grand_total_f = work_order_total + tender_premium_f
        grand_total_h = executed_total + tender_premium_h
        grand_total_j = overall_excess + tender_premium_j
        grand_total_l = overall_saving + tender_premium_l
        net_difference = grand_total_h - grand_total_f

        deviation_data["summary"] = {
            "work_order_total": round(work_order_total),
            "executed_total": round(executed_total),
            "overall_excess": round(overall_excess),
            "overall_saving": round(overall_saving),
            "premium": {"percent": premium_percent / 100, "type": premium_type},
            "tender_premium_f": tender_premium_f,
            "tender_premium_h": tender_premium_h,
            "tender_premium_j": tender_premium_j,
            "tender_premium_l": tender_premium_l,
            "grand_total_f": grand_total_f,
            "grand_total_h": grand_total_h,
            "grand_total_j": grand_total_j,
            "grand_total_l": grand_total_l,
            "net_difference": round(net_difference)
        }

    print(f"first_page_data['items'] type: {type(first_page_data['items'])}, length: {len(first_page_data['items'])}")
    print(f"extra_items_data['items'] type: {type(extra_items_data['items'])}, length: {len(extra_items_data['items'])}")
    if deviation_data:
        print(f"deviation_data['items'] type: {type(deviation_data['items'])}, length: {len(deviation_data['items'])}")
    
    return first_page_data, certificate_ii_data, certificate_iii_data, deviation_data, extra_items_data, note_sheet_data

def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount, note_sheet_data):
    percentage_work_done = float(payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    notes = []
    serial_number = 1
    notes.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1
    if percentage_work_done < 90:
        notes.append(f"{serial_number}. The execution of work at final stage is less than 90% of the Work Order Amount, the Requisite Deviation Statement is enclosed to observe check on unuseful expenditure. Approval of the Deviation is having jurisdiction under this office.")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        notes.append(f"{serial_number}. Requisite Deviation Statement is enclosed. The Overall Excess is less than or equal to 5% and is having approval jurisdiction under this office.")
        serial_number += 1
    elif percentage_work_done > 105:
        notes.append(f"{serial_number}. Requisite Deviation Statement is enclosed. The Overall Excess is more than 5% and Approval of the Deviation Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        serial_number += 1
    try:
        actual_dt = datetime.strptime(note_sheet_data["header"]["actual_completion"], '%d/%m/%Y')
        completion_dt = datetime.strptime(note_sheet_data["header"]["date_completion"], '%d/%m/%Y')
        delay_days = (actual_dt - completion_dt).days
        if delay_days > 0:
            time_allowed = (completion_dt - datetime.strptime(note_sheet_data["header"]["date_commencement"], '%d/%m/%Y')).days
            notes.append(f"{serial_number}. Time allowed for completion of the work was {time_allowed} days. The work was delayed by {delay_days} days.")
            serial_number += 1
            if delay_days > 0.5 * time_allowed:
                notes.append(f"{serial_number}. Approval of the Time Extension Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
            else:
                notes.append(f"{serial_number}. Approval of the Time Extension Case is to be done by this office.")
            serial_number += 1
        else:
            notes.append(f"{serial_number}. Work was completed in time.")
            serial_number += 1
    except (ValueError, TypeError):
        notes.append(f"{serial_number}. Unable to calculate delay due to invalid date format.")
        serial_number += 1
    if extra_item_amount > 0:
        extra_item_percentage = float(extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            notes.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; exceeds 5%, requires approval from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        else:
            notes.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; under 5%, approval of the same is to be granted by this office.")
        serial_number += 1
    notes.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1
    notes.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    notes.append("")
    notes.append("                                Premlata Jain")
    notes.append("                               AAO- As Auditor")
    return notes

def generate_pdf(sheet_name, data, orientation, output_path, note_sheet_data=None):
    st.write(f"Generating PDF for {sheet_name}, data type: {type(data)}")
    try:
        template_name = f"{sheet_name.lower().replace(' ', '_')}.html"
        template = env.get_template(template_name)
        context = {
            'data': data,
            'note_sheet_data': note_sheet_data if note_sheet_data else {},
            'header_data': data.get('header', {}) if sheet_name != "Note Sheet" else note_sheet_data.get('header', {}) if note_sheet_data else {}
        }
        # Log the note_sheet_data structure before rendering
        if sheet_name == "Note Sheet":
            print("note_sheet_data structure:", note_sheet_data)
            print("note_sheet_data keys:", list(note_sheet_data.keys()))
            print("work_order_amount in note_sheet_data:", note_sheet_data.get("work_order_amount", "Not found"))
        html_content = template.render(**context)
        options = {
            "page-size": "A4",
            "orientation": orientation,
            "margin-top": "0.25in" if sheet_name == "Note Sheet" else "0in",
            "margin-bottom": "0.6in" if sheet_name == "Note Sheet" else "0in",
            "margin-left": "0.25in" if sheet_name == "Note Sheet" else "0in",
            "margin-right": "0.25in" if sheet_name == "Note Sheet" else "0in",
            "encoding": "UTF-8"
        }
        pdfkit.from_string(html_content, output_path, configuration=config, options=options)
        st.write(f"Finished PDF for {sheet_name}")
    except TemplateNotFound:
        st.error(f"Template {template_name} not found in the templates directory.")
        raise
    except Exception as e:
        st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
        raise

def create_word_doc(sheet_name, data, doc_path, header_data=None):
    st.write(f"Creating Word doc for {sheet_name}")
    try:
        doc = Document()
        if sheet_name == "First Page":
            if header_data:
                for key, value in header_data.items():
                    doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            table = doc.add_table(rows=len(data["items"]) + 3, cols=8)
            table.style = "Table Grid"
            headers = ["Serial No.", "Unit", "Quantity", "Description", "Rate", "Amount", "BSR", "Remark"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("unit", ""))
                row.cells[2].text = str(item.get("quantity", ""))
                row.cells[3].text = str(item.get("description", ""))
                row.cells[4].text = str(item.get("rate", ""))
                row.cells[5].text = str(item.get("amount", ""))
                row.cells[6].text = str(item.get("bsr", ""))
                row.cells[7].text = str(item.get("remark", ""))
            row = table.rows[-3]
            row.cells[3].text = "Grand Total"
            row.cells[5].text = str(data["totals"].get("grand_total", ""))
            row = table.rows[-2]
            row.cells[3].text = f"Tender Premium @ {data['totals']['premium'].get('percent', 0) * 100:.2f}%"
            row.cells[5].text = str(data["totals"]["premium"].get("amount", ""))
            row = table.rows[-1]
            row.cells[3].text = "Payable Amount"
            row.cells[5].text = str(data["totals"].get("payable", ""))
        elif sheet_name == "Certificate II" or sheet_name == "Certificate III":
            doc.add_paragraph(f"Payable Amount: {data.get('payable_amount', '')}")
            doc.add_paragraph(f"Total in Words: {data.get('amount_words', '')}")
            if sheet_name == "Certificate III":
                doc.add_paragraph(data.get('certification', ''))
        elif sheet_name == "Extra Items":
            table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
            table.style = "Table Grid"
            headers = ["Serial No.", "Ref BSR", "Description", "Quantity", "Rate", "Amount", "Remark"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("ref_bsr", ""))
                row.cells[2].text = str(item.get("description", ""))
                row.cells[3].text = str(item.get("quantity", ""))
                row.cells[4].text = str(item.get("rate", ""))
                row.cells[5].text = str(item.get("amount", ""))
                row.cells[6].text = str(item.get("remark", ""))
        elif sheet_name == "Deviation Statement":
            if header_data:
                for key, value in header_data.items():
                    doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            table = doc.add_table(rows=len(data["items"]) + 5, cols=12)
            table.style = "Table Grid"
            headers = ["Serial No.", "Description", "Unit", "Qty WO", "Rate", "Amt WO", "Qty Bill", "Amt Bill", "Excess Qty", "Excess Amt", "Saving Qty", "Saving Amt"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("description", ""))
                row.cells[2].text = str(item.get("unit", ""))
                row.cells[3].text = str(item.get("qty_wo", ""))
                row.cells[4].text = str(item.get("rate", ""))
                row.cells[5].text = str(item.get("amt_wo", ""))
                row.cells[6].text = str(item.get("qty_bill", ""))
                row.cells[7].text = str(item.get("amt_bill", ""))
                row.cells[8].text = str(item.get("excess_qty", ""))
                row.cells[9].text = str(item.get("excess_amt", ""))
                row.cells[10].text = str(item.get("saving_qty", ""))
                row.cells[11].text = str(item.get("saving_amt", ""))
            row = table.rows[-4]
            row.cells[1].text = "Grand Total"
            row.cells[5].text = str(data["summary"].get("work_order_total", ""))
            row.cells[7].text = str(data["summary"].get("executed_total", ""))
            row.cells[9].text = str(data["summary"].get("overall_excess", ""))
            row.cells[11].text = str(data["summary"].get("overall_saving", ""))
            row = table.rows[-3]
            row.cells[1].text = f"Add Tender Premium @ {data['summary']['premium'].get('percent', 0) * 100:.2f}%"
            row.cells[5].text = str(data["summary"].get("tender_premium_f", ""))
            row.cells[7].text = str(data["summary"].get("tender_premium_h", ""))
            row.cells[9].text = str(data["summary"].get("tender_premium_j", ""))
            row.cells[11].text = str(data["summary"].get("tender_premium_l", ""))
            row = table.rows[-2]
            row.cells[1].text = "Grand Total including Tender Premium"
            row.cells[5].text = str(data["summary"].get("grand_total_f", ""))
            row.cells[7].text = str(data["summary"].get("grand_total_h", ""))
            row.cells[9].text = str(data["summary"].get("grand_total_j", ""))
            row.cells[11].text = str(data["summary"].get("grand_total_l", ""))
            row = table.rows[-1]
            net_difference = data["summary"].get("net_difference", 0)
            row.cells[1].text = "Overall Excess" if net_difference > 0 else "Overall Saving"
            row.cells[7].text = str(abs(round(net_difference)))
        elif sheet_name == "Note Sheet":
            if header_data:
                for key, value in header_data.items():
                    doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            for note in data.get("notes", []):
                doc.add_paragraph(str(note))
        doc.save(doc_path)
        st.write(f"Finished Word doc for {sheet_name}")
    except Exception as e:
        st.error(f"Error creating Word doc for {sheet_name}: {str(e)}")
        raise

# Streamlit app
st.title("Bill Generator")

# Sky blue gradient button CSS
st.markdown("""
    <style>
    .stButton > button {
        background: linear-gradient(to right, #87CEEB, #00B7EB);
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        cursor: pointer;
    }
    .stButton > button:hover {
        background: linear-gradient(to right, #00B7EB, #87CEEB);
    }
    </style>
""", unsafe_allow_html=True)

# Excel structure guidance
st.markdown("""
### Expected Excel Structure
Your Excel file must have three sheets: **"Work Order"**, **"Bill Quantity"**, and **"Extra Items"**. The sheets do not require specific column headers, as data is processed by row and column indices:

- **Work Order Sheet**:
  - Rows 1-20: Ignored (header data is taken from user inputs).
  - Row 21: Header row (e.g., 'Item', 'Unit', 'Quantity', 'Rate', 'Amount', 'BSR').
  - Rows 22+ (Excel row numbers):
    - Column A (index 0): Description
    - Column B (index 1): Unit
    - Column C (index 2): Quantity (numeric)
    - Column D (index 3): Rate (numeric)
    - Column E (index 4): Amount (calculated as Quantity * Rate)
    - Column F (index 5): BSR
    - Column G (index 6): Remark (optional)

- **Bill Quantity Sheet**:
  - Row 21: Header row.
  - Rows 22+:
    - Column A (index 0): Description
    - Column B (index 1): Unit
    - Column C (index 2): Quantity (numeric)
    - Column D (index 3): Rate (numeric)
    - Column E (index 4): Amount
    - Column F (index 5): BSR
    - Column G (index 6): Remark (optional)

- **Extra Items Sheet** (Optional):
  - If no extra items, the sheet can be empty or have any number of rows (no minimum required).
  - If extra items exist:
    - Rows 1-6: Ignored.
    - Row 7: Header row (e.g., 'S.No.', 'Ref. BSR No.', 'Particulars', 'Qty.', 'Rate', 'Amount').
    - Rows 7+:
      - Column A (index 0): Serial No.
      - Column B (index 1): Ref BSR
      - Column C (index 2): Description
      - Column D (index 3): Quantity (numeric)
      - Column E (index 4): Rate (numeric)
      - Column F (index 5): Amount
      - Column G (index 6): Remark (optional)

**Important Notes**:
- Ensure sheets are named exactly **"Work Order"**, **"Bill Quantity"**, and **"Extra Items"**.
- Data must start at row 22 for Work Order/Bill Quantity, row 7 for Extra Items (if present).
- Quantity and Rate columns must contain numeric values or properly formatted strings (e.g., '100', '50.25').
- Extra Items sheet does not require a Unit column; it will be left empty.
""")

with st.form(key='bill_form'):
    uploaded_files = st.file_uploader("Choose Excel file(s)", type="xlsx", accept_multiple_files=True)
    premium_percent = st.number_input("Tender Premium %", min_value=0.0, max_value=100.0, step=0.01, value=4.0)
    premium_type = st.selectbox("Premium Type", ["Above", "Below"], index=0)
    is_first_bill = st.selectbox("Is this the first bill?", ["Yes", "No"], index=0) == "Yes"
    is_final_bill = st.selectbox("Is this the final bill?", ["Yes", "No"], index=0) == "Yes"
    
    if not is_first_bill:
        st.markdown("**Please enter the amount paid in the last bill (required for non-first bills):**")
        amount_paid_last_bill = st.number_input("Amount Paid Vide Last Bill (Rs.)", min_value=0.0, step=1.0, value=0.0)
    else:
        amount_paid_last_bill = 0.0
        st.markdown("*Amount Paid Vide Last Bill is set to 0 for the first bill.*")
    
    st.subheader("Bill Information")
    user_inputs = {}
    user_inputs["work_order_amount"] = st.number_input("Work Order Amount (₹)", min_value=0.0, value=854678.0)
    user_inputs["date_commencement"] = st.text_input("Date of Written Order to Commence Work (DD/MM/YYYY)", value="09/01/2025")
    user_inputs["date_completion"] = st.text_input("Scheduled Date of Completion (DD/MM/YYYY)", value="17/04/2025")
    user_inputs["actual_completion"] = st.text_input("Date of Actual Completion (DD/MM/YYYY)", value="01/03/2025")
    user_inputs["name_of_firm"] = st.text_input("Name of Contractor/Supplier", value="M/s Seema Electrical, Udaipur")
    user_inputs["name_of_work"] = st.text_input("Name of Work", value="Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur")
    user_inputs["serial_no_bill"] = st.text_input("Serial No. of this Bill", value="First & Final Bill")
    user_inputs["work_order_ref"] = st.text_input("Reference to Work Order/Agreement", value="1179 dated 09-01-2025")
    user_inputs["agreement_no"] = st.text_input("Agreement No.", value="48/2024-25")
    user_inputs["measurement_date"] = st.text_input("Measurement Date (DD/MM/YYYY)", value="03/03/2025")

    submitted = st.form_submit_button("Generate Bill")

if submitted and uploaded_files:
    try:
        TEMP_DIR = tempfile.mkdtemp()
        zip_path = os.path.join(TEMP_DIR, f"BILL_OUTPUT_{datetime.now().strftime('%Y%m%d')}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for uploaded_file in uploaded_files:
                excel_file = pd.ExcelFile(uploaded_file)
                # Verify sheet names
                expected_sheets = ["Work Order", "Bill Quantity", "Extra Items"]
                if not all(sheet in excel_file.sheet_names for sheet in expected_sheets):
                    missing_sheets = [sheet for sheet in expected_sheets if sheet not in excel_file.sheet_names]
                    raise ValueError(f"Excel file missing required sheets: {missing_sheets}")

                ws_wo = excel_file.parse("Work Order", header=None)
                ws_bq = excel_file.parse("Bill Quantity", header=None)
                ws_extra = excel_file.parse("Extra Items", header=None)

                first_page_data, certificate_ii_data, certificate_iii_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
                    ws_wo, ws_bq, ws_extra, premium_percent, premium_type.lower(),
                    amount_paid_last_bill, is_first_bill, is_final_bill, user_inputs
                )

                # Log note_sheet_data after processing
                print("note_sheet_data after process_bill:", note_sheet_data)

                # Generate note sheet notes
                note_sheet_data["notes"] = generate_bill_notes(
                    first_page_data["totals"]["payable"],
                    note_sheet_data["work_order_amount"],  # Use directly from note_sheet_data
                    extra_items_data["totals"]["payable"],
                    note_sheet_data
                )

                pdf_files = []
                word_files = []

                # Generate PDFs
                pdf_sheet_names = [
                    ("First Page", first_page_data, "portrait"),
                    ("Certificate II", certificate_ii_data, "portrait"),
                    ("Certificate III", certificate_iii_data, "portrait")
                ]
                if extra_items_data["items"]:
                    pdf_sheet_names.append(("Extra Items", extra_items_data, "portrait"))
                if is_final_bill and deviation_data:
                    pdf_sheet_names.append(("Deviation Statement", deviation_data, "landscape"))
                if is_final_bill:
                    pdf_sheet_names.append(("Note Sheet", note_sheet_data, "portrait"))

                for sheet_name, data, orientation in pdf_sheet_names:
                    pdf_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}_{uploaded_file.name}.pdf")
                    generate_pdf(sheet_name, data, orientation, pdf_path, note_sheet_data if sheet_name == "Note Sheet" else None)
                    pdf_files.append(pdf_path)

                # Generate Word documents
                word_sheet_names = ["First Page", "Certificate II", "Certificate III"]
                if extra_items_data["items"]:
                    word_sheet_names.append("Extra Items")
                if is_final_bill and deviation_data:
                    word_sheet_names.append("Deviation Statement")
                if is_final_bill:
                    word_sheet_names.append("Note Sheet")

                for sheet_name in word_sheet_names:
                    data = {
                        "First Page": first_page_data,
                        "Certificate II": certificate_ii_data,
                        "Certificate III": certificate_iii_data,
                        "Extra Items": extra_items_data,
                        "Deviation Statement": deviation_data,
                        "Note Sheet": note_sheet_data
                    }[sheet_name]
                    doc_path = os.path.join(TEMP_DIR, f"{sheet_name.replace(' ', '_')}_{uploaded_file.name}.docx")
                    create_word_doc(sheet_name, data, doc_path, first_page_data["header"])
                    word_files.append(doc_path)

                # Combine PDFs
                pdf_output = os.path.join(TEMP_DIR, f"BILL_AND_DEVIATION_{datetime.now().strftime('%Y%m%d')}_{uploaded_file.name}.pdf")
                writer = PdfWriter()
                for pdf in pdf_files:
                    if os.path.exists(pdf):
                        reader = PdfReader(pdf)
                        for page in reader.pages:
                            writer.add_page(page)
                with open(pdf_output, "wb") as out_file:
                    writer.write(out_file)
                zipf.write(pdf_output, os.path.basename(pdf_output))
                for file in pdf_files + word_files:
                    if os.path.exists(file):
                        zipf.write(file, os.path.basename(file))

        with open(zip_path, "rb") as f:
            st.download_button(
                label="Download Bill Output",
                data=f,
                file_name="bill_output.zip",
                mime="application/zip"
            )
    except Exception as e:
        st.error(f"Error: {str(e)}")
        st.write(traceback.format_exc())
    finally:
        if os.path.exists(TEMP_DIR):
            shutil.rmtree(TEMP_DIR)