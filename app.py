import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, date
import os
import webbrowser
import time
import tempfile
import base64
from jinja2 import Environment, FileSystemLoader
from utils import (
    process_bill,
    generate_bill_notes,
    generate_pdf,
    create_word_doc,
    number_to_words
)
import traceback
import shutil
import subprocess
import zipfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)
env.filters['strptime'] = lambda s, fmt: datetime.strptime(s, fmt) if s else None

st.title("Contractor Bill Generator ")

# Add custom CSS for styling
st.markdown("""
<style>
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 4px;
        cursor: pointer;
        font-size: 16px;
    }
    .stButton > button:hover {
        background-color: #45a049;
    }
    .stFileUploader > div > div {
        background-color: #f5f5f5;
        padding: 10px;
        border-radius: 4px;
        border: 1px solid #ddd;
    }
    .stFileUploader > div > div:hover {
        background-color: #e0e0e0;
    }
    .stFormSubmitButton > button {
        background-color: #2196F3 !important;
        color: white !important;
    }
    .stFormSubmitButton > button:hover {
        background-color: #1976D2 !important;
    }
</style>
""", unsafe_allow_html=True)

# Application description
st.markdown("""
<div style='text-align: center; padding: 20px; background-color: #f9f9f9; border-radius: 8px;'>
    <h3>Generate Contractor Bills with Ease</h3>
    <p>Upload your Excel files and generate professional contractor bills in seconds.</p>
</div>
""", unsafe_allow_html=True)

# Main content
st.markdown("""
### Instructions:
1. Fill in the required details in the sidebar
2. Upload an Excel file containing three sheets:
   - Work Order (ws_wo)
   - Bill Quantity (ws_bq)
   - Extra Items (ws_extra)
3. View the bill summary and download the processed documents
""")

# Main content area for mandatory fields
st.markdown("""
<div style='display: flex; justify-content: space-between; gap: 20px;'>
    <div style='flex: 1; border: 1px solid #ddd; padding: 20px; border-radius: 8px; background-color: #f8f9fa;'>
        <h4>Optional Fields</h4>
        <div style='margin-top: 20px;'>
            <label>Contractor Name</label>
            <input type="text" class="stTextInput" id="contractor_name">
            <label>Work Name</label>
            <input type="text" class="stTextInput" id="work_name">
            <label>Bill Serial Number</label>
            <input type="text" class="stTextInput" id="bill_serial">
            <label>Agreement Number</label>
            <input type="text" class="stTextInput" id="agreement_no">
            <label>Work Order Reference</label>
            <input type="text" class="stTextInput" id="work_order_ref">
        </div>
    </div>
    <div style='flex: 1; border: 1px solid #ddd; padding: 20px; border-radius: 8px;'>
        <h4>Mandatory Fields</h4>
        <div style='margin-top: 20px;'>
            <label>Start Date *</label>
            <input type="date" class="stDateInput" id="start_date">
            
            <label>Scheduled Completion Date *</label>
            <input type="date" class="stDateInput" id="completion_date">
            
            <label>Actual Completion Date *</label>
            <input type="date" class="stDateInput" id="actual_completion_date">
            
            <label>Order Date *</label>
            <input type="date" class="stDateInput" id="order_date">
            
            <label>Measurement Date</label>
            <input type="date" class="stDateInput" id="measurement_date">
            
            <label>Bill Type *</label>
            <select class="stSelectbox" id="bill_type">
                <option value="">Select Bill Type</option>
                <option value="final">Final Bill</option>
                <option value="running">Running Bill</option>
            </select>
            
            <label>Premium</label>
            <div style='display: flex; gap: 10px;'>
                <input type="number" class="stNumberInput" id="premium_amount">
                <select class="stSelectbox" id="premium_type">
                    <option value="above">Above</option>
                    <option value="below">Below</option>
                </select>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# File upload section
st.markdown("""
<div style='border: 1px solid #ddd; padding: 20px; border-radius: 8px; margin: 20px 0;'>
    <h4>Upload Excel File</h4>
    <p style='color: #666;'>Select an Excel file containing the required sheets</p>
    <div class="stFileUploader">
        <div>
            <label for="file_uploader">Choose file</label>
            <input type="file" id="file_uploader" accept=".xlsx,.xls">
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Sidebar for input parameters
st.sidebar.markdown("""
<style>
    .sidebar .sidebar-content {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 8px;
    }
    .stForm {
        background-color: white;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
</style>
""", unsafe_allow_html=True)

st.sidebar.header("Bill Information")

# Required fields with validation
with st.sidebar.form("bill_info_form", clear_on_submit=False):
    st.subheader("Mandatory Fields")
    
    # Bold labels for mandatory fields
    st.markdown("<strong style='color:red'>* Required fields</strong>", unsafe_allow_html=True)
    
    # Date fields
    col1, col2 = st.columns(2)
    with col1:
        start_date = st.date_input(
            "Start Date *",
            date.today(),
            help="The date when work started",
            key="sidebar_start_date"
        )
        
        completion_date = st.date_input(
            "Scheduled Completion Date *",
            date.today(),
            help="The date when work was scheduled to complete",
            key="sidebar_completion_date"
        )
        
        # Measurement Date is non-mandatory
        measurement_date = st.date_input(
            "Measurement Date",
            None,
            help="The date when measurements were taken (optional)",
            key="sidebar_measurement_date"
        )
    
    with col2:
        actual_completion_date = st.date_input(
            "Actual Completion Date *",
            date.today(),
            help="The actual date when work was completed",
            key="sidebar_actual_completion_date"
        )
        
        order_date = st.date_input(
            "Order Date *",
            date.today(),
            help="Date of written order to commence work",
            key="sidebar_order_date"
        )
    
    # Contractor and work details (all optional)
    st.markdown("---")
    st.subheader("Contractor & Work Details")
    
    contractor_name = st.text_input(
        "Contractor Name",
        "",
        help="Name of the contractor (optional)",
        key="sidebar_contractor_name"
    )
    
    work_name = st.text_input(
        "Work Name",
        "",
        help="Name of the work/project (optional)",
        key="sidebar_work_name"
    )
    
    bill_serial = st.text_input(
        "Bill Serial Number",
        "",
        help="Serial number of this bill (optional)",
        key="sidebar_bill_serial"
    )
    
    agreement_no = st.text_input(
        "Agreement Number",
        "",
        help="Number of the agreement (optional)",
        key="sidebar_agreement_no"
    )
    
    work_order_ref = st.text_input(
        "Work Order Reference",
        "",
        help="Reference number of the work order (optional)",
        key="sidebar_work_order_ref"
    )
    
    # Financial details
    st.markdown("---")
    st.subheader("Financial Details")
    
    work_order_amount = st.number_input(
        "Work Order Amount *",
        min_value=0.0,
        help="Total amount in the work order",
        key="sidebar_work_order_amount"
    )
    
    # Premium details (always fixed)
    premium_percent = st.number_input(
        "Premium Percentage",
        min_value=-100.0,
        max_value=100.0,
        value=0.0,
        help="Premium percentage (positive for above, negative for below)",
        key="sidebar_premium_percent"
    )
    premium_type = "Fixed"  # Always fixed
    
    # Bill type and number
    bill_type = st.selectbox(
        "Bill Type",
        ["Running Bill", "Final Bill"],
        help="Type of bill being generated",
        key="sidebar_bill_type_select"
    )
    
    bill_number = st.selectbox(
        "Bill Number",
        ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth", "Tenth"],
        help="The sequence number of this bill in the series",
        key="sidebar_bill_number_select"
    )
    
    # Amount details
    if bill_number != "First":
        amount_paid_last_bill = st.number_input(
            "Amount Paid in Last Bill",
            min_value=0.0,
            help="Amount paid in the previous bill",
            key="sidebar_amount_paid_last_bill"
        )
    else:
        amount_paid_last_bill = 0.0  # First bill has no previous payment
    
    # Non-mandatory fields
    st.markdown("---")
    st.subheader("Optional Fields")
    
    # Last Bill Reference
    last_bill_reference = st.text_input(
        "Last Bill Reference",
        "",
        help="Reference to the previous bill (optional)",
        key="sidebar_last_bill_reference"
    )
    
    # Measurement Date
    measurement_date = st.date_input(
        "Measurement Date",
        None,
        help="The date when measurements were taken (optional)",
        key="sidebar_measurement_date_2"
    )
    
    # Bill type selection
    bill_type = st.radio(
        "Bill Type",
        ["Running Bill", "Final Bill"],
        help="Select whether this is a running bill or the final bill",
        key="sidebar_bill_type_radio"
    )
    
    # Bill number selection
    bill_number = st.selectbox(
        "Bill Number",
        ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth", "Tenth"],
        help="Select the sequential number of this bill",
        key="sidebar_bill_number_select_2"
    )
    
    # Last bill reference
    last_bill_options = ["Not Applicable"]
    if bill_number != "First":
        bill_index = ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth", "Tenth"].index(bill_number)
        previous_bills = [f"{prev} Running Bill" for prev in ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth"][:bill_index]]
        last_bill_options = previous_bills + last_bill_options
    
    last_bill = st.selectbox(
        "Last Bill Reference",
        last_bill_options,
        help="Reference to the previous bill if applicable",
        key="sidebar_last_bill_select"
    )
    
    submitted = st.form_submit_button("Submit Details")

# Check for required dependencies
import subprocess

def check_dependencies():
    def check_dependency(executable):
        try:
            subprocess.run([executable, "--version"], check=True, capture_output=True)
            return True
        except:
            return False

    # Check for required dependencies
    dependencies = {
        "wkhtmltopdf": check_dependency("wkhtmltopdf"),
        "pdftk": check_dependency("pdftk")
    }

    # Display dependency status
    st.sidebar.markdown("---")
    st.sidebar.header("System Status")
    for dep, status in dependencies.items():
        status_text = "Installed" if status else "Not Installed"
        st.sidebar.markdown(f"**{dep}:** {status_text}")

    # Process form submission
    if submitted and uploaded_file is not None:
        try:
            # Read the uploaded file
            with pd.ExcelFile(uploaded_file) as xls:
                ws_wo = pd.read_excel(xls, "Work Order", header=None)
                ws_bq = pd.read_excel(xls, "Bill Quantity", header=None)
                ws_extra = pd.read_excel(xls, "Extra Items", header=None)

            # Process the bill
            first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
                ws_wo, ws_bq, ws_extra,
                premium_percent=premium_percent,
                premium_type=premium_type,
                amount_paid_last_bill=amount_paid_last_bill,
                is_first_bill=bill_number == "First",
                user_inputs={
                    "start_date": start_date,
                    "completion_date": completion_date,
                    "actual_completion_date": actual_completion_date,
                    "measurement_date": measurement_date,
                    "order_date": order_date,
                    "contractor_name": contractor_name,
                    "work_name": work_name,
                    "bill_serial": bill_serial,
                    "agreement_no": agreement_no,
                    "work_order_ref": work_order_ref,
                    "work_order_amount": work_order_amount,
                    "bill_type": bill_type,
                    "bill_number": bill_number,
                    "last_bill": last_bill_reference
                }
            )

            # Generate PDF
            pdf_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 
                                   "output", 
                                   f"bill_{bill_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

            # Generate each page separately
            if first_page_data:
                first_page_data['totals'] = {
                    'work_order_total': first_page_data.get('work_order_total', 0),
                    'premium_amount': first_page_data.get('premium_amount', 0),
                    'grand_total': first_page_data.get('grand_total', 0)
                }
                generate_pdf("first_page", first_page_data, "portrait", pdf_path, True)

            if last_page_data:
                generate_pdf("certificate_ii" if bill_type == "Running Bill" else "certificate_iii", last_page_data, "portrait", pdf_path)

            if deviation_data:
                generate_pdf("deviation_statement", deviation_data, "portrait", pdf_path)

            if extra_items_data:
                extra_items_data['premium'] = {
                    'percent': extra_items_data.get('premium_percent', 0),
                    'amount': extra_items_data.get('premium_amount', 0)
                }
                generate_pdf("extra_items", extra_items_data, "portrait", pdf_path)

            if note_sheet_data:
                note_sheet_data['header'] = {
                    'agreement_no': note_sheet_data.get('agreement_no', ''),
                    'work_order_amount': note_sheet_data.get('work_order_amount', 0)
                }
                generate_pdf("note_sheet", note_sheet_data, "portrait", pdf_path)

            # Display success message and download link
            st.success("Bill processed successfully!")
            
            # Create download button for PDF
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name=f"bill_{bill_number}.pdf",
                mime="application/pdf",
                key="pdf_download"
            )

            # Display summary
            if first_page_data and 'totals' in first_page_data:
                st.markdown(f"""
                <div style='border: 1px solid #ddd; padding: 20px; border-radius: 8px; margin: 20px 0; background-color: #f8f9fa;'>
                    <h4>Bill Summary</h4>
                    <p><strong>Total Work Order Amount:</strong> ₹{first_page_data['totals']['work_order_total']:,}</p>
                    <p><strong>Premium Amount:</strong> ₹{first_page_data['totals']['premium_amount']:,}</p>
                    <p><strong>Grand Total:</strong> ₹{first_page_data['totals']['grand_total']:,}</p>
                </div>
                """, unsafe_allow_html=True)

        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            st.write(traceback.format_exc())
        try:
            subprocess.run([executable, "--version"], capture_output=True, check=True)
            return True
        except Exception as e:
            return False
    
    if not check_dependency("wkhtmltopdf"):
        st.error("wkhtmltopdf is not installed. Please install it from https://wkhtmltopdf.org/downloads.html")
        return False
    
    if not check_dependency("pdftk"):
        st.error("pdftk is not installed. Please install it from https://www.pdflabs.com/tools/pdftk-the-pdf-toolkit/")
        return False
    
    return True

def check_and_display_dependencies():
    try:
        if not check_dependencies():
            st.stop()  # Stop the Streamlit app if dependencies are missing
    except Exception as e:
        st.error(f"Error checking dependencies: {str(e)}")
        st.stop()  # Stop the Streamlit app if there's an error

check_and_display_dependencies()

# Initialize session state
if "form_data" not in st.session_state:
    st.session_state.form_data = {}
if "processing" not in st.session_state:
    st.session_state.processing = False
if "error" not in st.session_state:
    st.session_state.error = None

def read_excel_file(xls, sheet_name):
    """Read Excel sheet with proper column names"""
    # Read the sheet
    df = pd.read_excel(xls, sheet_name=sheet_name)
    
    # Debug print to check original data
    print(f"\nOriginal {sheet_name} sheet:")
    print(df.head())
    
    # Skip rows based on sheet name
    if sheet_name in ["Bill Quantity", "Work Order"]:
        df = df.iloc[19:]  # Skip first 19 rows
    elif sheet_name == "Extra Items":
        df = df.iloc[5:]   # Skip first 5 rows
    
    # Reset index
    df = df.reset_index(drop=True)
    
    # Debug print to check data after skipping rows
    print(f"\n{sheet_name} sheet after skipping rows:")
    print(df.head())
    
    return df

def generate_pdf(template_name, data, orientation, output_path, is_first_page=False):
    """
    Generate PDF from HTML template.
    
    Args:
        template_name: Name of the template file (without .html extension)
        data: Dictionary containing data for template
        orientation: Page orientation (portrait/landscape)
        output_path: Path where PDF should be saved
        is_first_page: Boolean indicating if this is the first page
    """
    try:
        # Validate template name
        if not template_name:
            raise ValueError("Template name cannot be empty")
        
        # Validate orientation
        if orientation not in ["portrait", "landscape"]:
            raise ValueError("Orientation must be either 'portrait' or 'landscape'")
        
        # Validate output path
        if not output_path:
            raise ValueError("Output path cannot be empty")
        
        # Load template
        template_path = os.path.join("templates", f"{template_name}.html")
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
            
        # Create temporary HTML file
        temp_html = os.path.join(os.path.dirname(output_path), f"{template_name}.html")
        
        # Load and render template using Jinja2
        try:
            env = Environment(loader=FileSystemLoader('templates'))
            template = env.get_template(f"{template_name}.html")
            rendered_html = template.render(data)
        except Exception as e:
            raise ValueError(f"Error rendering template {template_name}: {str(e)}")
            
        # Create temporary directory for PDFs
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Generate PDF for this template
            temp_pdf = os.path.join(temp_dir, f"{template_name}.pdf")
            
            # Write rendered HTML to file
            try:
                with open(temp_html, "w", encoding="utf-8") as f:
                    f.write(rendered_html)
            except Exception as e:
                raise IOError(f"Error writing HTML file: {str(e)}")
            
            # Generate PDF using wkhtmltopdf
            cmd = [
                "wkhtmltopdf",
                "--orientation", orientation,
                "--page-size", "A4",
                "--margin-top", "10mm",
                "--margin-bottom", "10mm",
                "--margin-left", "15mm",
                "--margin-right", "15mm",
                "--quiet",
                temp_html,
                temp_pdf
            ]
            
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            except subprocess.CalledProcessError as e:
                raise RuntimeError(f"Error generating PDF: {e.stderr}\nCommand: {' '.join(cmd)}")
            except FileNotFoundError:
                raise RuntimeError("wkhtmltopdf executable not found. Please install it from https://wkhtmltopdf.org/downloads.html")
            
            # Verify PDF was created
            if not os.path.exists(temp_pdf):
                raise IOError(f"Failed to create PDF for {template_name}")
                
            # Verify PDF size (should be at least 1KB)
            if os.path.getsize(temp_pdf) < 1024:
                raise IOError(f"Generated PDF for {template_name} is too small")
                
            # If this is the first page, create the final PDF
            if is_first_page:
                try:
                    shutil.copy2(temp_pdf, output_path)
                except Exception as e:
                    raise IOError(f"Error copying PDF: {str(e)}")
            else:
                # For subsequent pages, merge with existing PDF
                try:
                    cmd = [
                        "pdftk",
                        output_path,
                        temp_pdf,
                        "cat",
                        "output",
                        output_path
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                except subprocess.CalledProcessError as e:
                    raise RuntimeError(f"Error merging PDF pages: {e.stderr}\nCommand: {' '.join(cmd)}")
                except FileNotFoundError:
                    raise RuntimeError("pdftk executable not found. Please install it from https://www.pdflabs.com/tools/pdftk-the-pdf-toolkit/")
            
            return True
            
        finally:
            # Clean up temporary files
            if os.path.exists(temp_html):
                os.remove(temp_html)
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
                
    except Exception as e:
        print(f"Error in generate_pdf: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False
        # Load template
        template_path = os.path.join("templates", f"{template_name}.html")
        if not os.path.exists(template_path):
            st.error(f"Template file not found: {template_path}")
            return False
            
        # Create temporary HTML file
        temp_html = os.path.join(os.path.dirname(output_path), f"{template_name}.html")
        
        # Load and render template using Jinja2
        env = Environment(loader=FileSystemLoader('templates'))
        template = env.get_template(f"{template_name}.html")
        rendered_html = template.render(data)
        
        # Create temporary directory for PDFs
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Generate PDF for this template
            temp_pdf = os.path.join(temp_dir, f"{template_name}.pdf")
            
            # Write rendered HTML to file
            with open(temp_html, "w", encoding="utf-8") as f:
                f.write(rendered_html)
            
            # Generate PDF using wkhtmltopdf
            cmd = [
                "wkhtmltopdf",
                "--orientation", orientation,
                "--page-size", "A4",
                "--margin-top", "10mm",
                "--margin-bottom", "10mm",
                "--margin-left", "15mm",
                "--margin-right", "15mm",
                "--quiet",
                temp_html,
                temp_pdf
            ]
            
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            except subprocess.CalledProcessError as e:
                st.error(f"Error generating PDF for {template_name}: {e.stderr}")
                st.error(f"Command: {' '.join(cmd)}")
                return False
            except FileNotFoundError:
                st.error("wkhtmltopdf executable not found. Please install it from https://wkhtmltopdf.org/downloads.html")
                return False
            
            # Verify PDF was created
            if not os.path.exists(temp_pdf):
                st.error(f"Failed to create PDF for {template_name}")
                return False
                
            # Verify PDF size (should be at least 1KB)
            if os.path.getsize(temp_pdf) < 1024:
                st.error(f"Generated PDF for {template_name} is too small")
                return False
                
            # If this is the first page, create the final PDF
            if is_first_page:
                try:
                    shutil.copy2(temp_pdf, output_path)
                except Exception as e:
                    st.error(f"Error copying PDF: {str(e)}")
                    return False
            else:
                # For subsequent pages, merge with existing PDF
                try:
                    cmd = [
                        "pdftk",
                        output_path,
                        temp_pdf,
                        "cat",
                        "output",
                        output_path
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
                except subprocess.CalledProcessError as e:
                    st.error(f"Error merging PDF pages: {e.stderr}")
                    st.error(f"Command: {' '.join(cmd)}")
                    return False
                except FileNotFoundError:
                    st.error("pdftk executable not found. Please install it from https://www.pdflabs.com/tools/pdftk-the-pdf-toolkit/")
                    return False
                    
            # Verify final PDF size
            if os.path.getsize(output_path) < 1024:
                st.error("Final PDF is too small")
                return False
                
            return True
                
        finally:
            # Clean up temporary files
            if os.path.exists(temp_html):
                os.remove(temp_html)
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            
    except Exception as e:
        st.error(f"Error in generate_pdf: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

def create_word_doc(template_name, data, output_path):
    """
    Create Word document from template.
    
    Args:
        template_name: Name of the template file (without .docx extension)
        data: Dictionary containing data for template
        output_path: Path where Word document should be saved
    """
    try:
        # Create a new Word document
        doc = Document()
        
        # Add heading
        heading = doc.add_heading('CONTRACTOR BILL', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add main content based on template
        if template_name == "first_page":
            # Add header table
            header_table = doc.add_table(rows=len(data["header"]), cols=7)
        elif template_name == "certificate_ii":
            # Add certificate II text
            doc.add_heading("Certificate II", 0)
            doc.add_paragraph("This is to certify that...")

        elif template_name == "certificate_iii":
            # Add certificate III text
            doc.add_heading("Certificate III", 0)
            doc.add_paragraph("This is to certify that...")

        elif template_name == "deviation_statement":
            # Add deviation statement text
            doc.add_heading("Deviation Statement", 0)
            doc.add_paragraph("The following deviations have been made...")

            # Add deviations table
            if 'deviations' in data:
                deviations_table = doc.add_table(rows=1, cols=4)
                deviations_table.style = 'Table Grid'

                # Add table headers
                headers = ["S.No.", "Item", "Deviation", "Reason"]
                for i, header in enumerate(headers):
                    cell = deviations_table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True

                # Add deviations
                for idx, deviation in enumerate(data['deviations'], 1):
                    row = deviations_table.add_row()
                    for i, key in enumerate(["sno", "item", "deviation", "reason"]):
                        cell = row.cells[i]
                        cell.text = str(deviation.get(key, ""))

        elif template_name == "note_sheet":
            # Add note sheet text
            doc.add_heading("NOTE SHEET", 0)
            doc.add_paragraph("\n")

            # Add notes
            for note in data["notes"]:
                doc.add_paragraph(note)

        elif template_name == "extra_items":
            # Add extra items text
            doc.add_heading("EXTRA ITEMS", 0)
            doc.add_paragraph("\n")

            # Add items table
            items_table = doc.add_table(rows=1, cols=8)
            items_table.style = 'Table Grid'

            # Add table headers
            headers = ["Serial No.", "Description", "Unit", "Quantity", "Rate", "Amount", "Remark"]
            for i, header in enumerate(headers):
                cell = items_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            # Add items
            for item in data["items"]:
                row = items_table.add_row()
                for i, key in enumerate(["serial_no", "description", "unit", "quantity", "rate", "amount", "remark"]):
                    cell = row.cells[i]
                    cell.text = str(item.get(key, ""))

        # Save document
        doc.save(output_path)

        # Verify file was created
        if not os.path.exists(output_path):
            st.error(f"Failed to create Word document: {output_path}")
            return False

        # Verify file size (should be at least 1KB)
        if os.path.getsize(output_path) < 1024:
            st.error(f"Generated Word document is too small: {output_path}")
            return False

        return True

    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

def handle_form_submission(uploaded_file):
    """
    Handle form submission and process bill generation.
    
    Args:
        uploaded_file: Uploaded Excel file
    """
    try:
        # Validate form data
        if not all([st.session_state.form_data.get(f) for f in ["start_date", "completion_date", "bill_type"]]):
            raise ValueError("Please fill in all required fields")
            
        if st.session_state.form_data["start_date"] > st.session_state.form_data["completion_date"]:
            raise ValueError("Start date cannot be after completion date")
            
        # Validate file
        if not uploaded_file:
            raise ValueError("Please upload an Excel file")
            
        # Show loading state
        st.session_state.processing = True
        st.session_state.error = None
        
        # Process file with context manager
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            
            # Save uploaded file to temp directory
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            # Read Excel file with optimized memory usage
            with pd.ExcelFile(temp_path, engine='openpyxl') as xls:
                # Validate sheet names
                required_sheets = ["Work Order", "Bill Quantity", "Extra Items"]
                missing_sheets = [sheet for sheet in required_sheets if sheet not in xls.sheet_names]
                if missing_sheets:
                    raise ValueError(f"Missing required sheets: {', '.join(missing_sheets)}")
                
                # Process each sheet
                ws_bq = read_excel_file(xls, "Bill Quantity")
                ws_wo = read_excel_file(xls, "Work Order")
                ws_extra = read_excel_file(xls, "Extra Items")
                
                # Get form data with type validation
                bill_type = st.session_state.form_data.get("bill_type", "").lower()
                if bill_type not in ["final bill", "running bill"]:
                    raise ValueError("Invalid bill type. Must be either 'Final Bill' or 'Running Bill'")
                    
                premium_percent = float(st.session_state.form_data.get("premium_amount", 0))
                if not (0 <= premium_percent <= 100):
                    raise ValueError("Premium percentage must be between 0 and 100")
                    
                premium_type = st.session_state.form_data.get("premium_type", "above")
                amount_paid_last_bill = float(st.session_state.form_data.get("amount_paid_last_bill", 0))
                is_first_bill = st.session_state.form_data["bill_number"] == "First"
                
                # Process bill with error handling
                try:
                    result = process_bill(
                        ws_wo, ws_bq, ws_extra,
                        premium_percent, premium_type,
                        amount_paid_last_bill, is_first_bill,
                        st.session_state.form_data
                    )
                except ValueError as e:
                    st.error(f"Error processing bill: {str(e)}")
                    return
                
                # Generate documents in parallel
                from concurrent.futures import ThreadPoolExecutor
                
                with ThreadPoolExecutor(max_workers=2) as executor:
                    # Generate PDFs
                    pdf_futures = [
                        executor.submit(generate_pdf, "first_page", result[0], "portrait", os.path.join(temp_dir, "first_page.pdf")),
                        executor.submit(generate_pdf, "bill_summary", result[1], "portrait", os.path.join(temp_dir, "bill_summary.pdf"))
                    ]
                    
                    # Wait for all PDFs to complete
                    for future in pdf_futures:
                        future.result()
                
                # Create zip file with all documents
                zip_path = os.path.join(temp_dir, "bill_documents.zip")
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            if file.endswith(('.pdf', '.docx')):
                                file_path = os.path.join(root, file)
                                zipf.write(file_path, os.path.basename(file_path))
                
                # Provide download link
                with open(zip_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                    href = f'<a href="data:application/octet-stream;base64,{b64}" download="bill_documents.zip">Download Documents</a>'
                    st.markdown(href, unsafe_allow_html=True)
                    
                # Display success message
                st.session_state.processing = False
                st.success("Bill processing completed successfully!")
                
                # Display summary
                if result[0] and 'totals' in result[0]:
                    st.write("""
                        <div style='border: 1px solid #ddd; padding: 20px; border-radius: 8px; margin: 20px 0; background-color: #f8f9fa;'>
                            <h4>Bill Summary</h4>
                            <p><strong>Total Work Order Amount:</strong> ₹{result[0]['totals']['work_order_total']:,}</p>
                            <p><strong>Premium Amount:</strong> ₹{result[0]['totals']['premium_amount']:,}</p>
                            <p><strong>Grand Total:</strong> ₹{result[0]['totals']['grand_total']:,}</p>
                        </div>
                        """, unsafe_allow_html=True)

    except Exception as e:
        error_msg = f"Error processing bill: {str(e)}"
        st.session_state.processing = False
        st.session_state.error = error_msg
        st.error(error_msg)
        st.write(traceback.format_exc())

        # Log the error for debugging
        with open("error_log.txt", "a") as f:
            f.write(f"\n{datetime.now()}: {error_msg}\n")
            f.write(traceback.format_exc())
            f.write("\n" * 3)

# Main content
with st.form("bill_form", clear_on_submit=False):
    # Get form data
    for field in ["start_date", "completion_date", "actual_completion_date", "order_date"]:
        st.session_state.form_data[field] = st.date_input(
            f"{field.replace('_', ' ').title()} *",
            st.session_state.form_data.get(field, date.today()),
            help=f"{field.replace('_', ' ').title()} date",
            key=f"{field}_input"
        )
        
    st.session_state.form_data["bill_type"] = st.radio(
        "Bill Type",
        ["Running Bill", "Final Bill"],
        index=0 if st.session_state.form_data.get("bill_type") == "Running Bill" else 1,
        help="Select whether this is a running bill or the final bill",
        key="bill_type_radio"
    )
    
    st.session_state.form_data["bill_number"] = st.selectbox(
        "Bill Number",
        ["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth", "Tenth"],
        index=["First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh", "Eighth", "Ninth", "Tenth"].index(st.session_state.form_data.get("bill_number", "First")),
        help="Select the sequential number of this bill",
        key="bill_number_select"
    )
    
    # Premium section
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.form_data["premium_amount"] = st.number_input(
            "Premium Amount",
            min_value=0,
            max_value=100,
            value=st.session_state.form_data.get("premium_amount", 0),
            help="Enter premium percentage",
            key="premium_amount_input"
        )
    with col2:
        st.session_state.form_data["premium_type"] = st.selectbox(
            "Premium Type",
            ["Above", "Below"],
            index=["Above", "Below"].index(st.session_state.form_data.get("premium_type", "Above")),
            help="Select whether premium is above or below",
            key="premium_type_select"
        )
    
    # Optional fields
    st.session_state.form_data["contractor_name"] = st.text_input(
        "Contractor Name",
        st.session_state.form_data.get("contractor_name", ""),
        help="Enter contractor's name",
        key="contractor_name_input"
    )
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload Excel File",
        type=["xlsx"],
        help="Upload Excel file containing Work Order, Bill Quantity, and Extra Items sheets",
        key="file_uploader"
    )
    
    # Submit button
    submit_button = st.form_submit_button("Generate Bill")
    
    # Process form submission
    if submit_button:
        if not uploaded_file:
            st.error("Please upload an Excel file")
        else:
            try:
                handle_form_submission(uploaded_file)
            except Exception as e:
                st.error(f"Error processing bill: {str(e)}")
                st.write(traceback.format_exc())
    
    # Show loading state
    if st.session_state.processing:
        with st.spinner("Processing bill..."):
            pass
    
    # Show error if present
    if st.session_state.error:
        st.error(st.session_state.error)

# Add clear form button
if st.button("Clear Form"):
    st.session_state.form_data = {}
    st.session_state.processing = False
    st.session_state.error = None
    st.experimental_rerun()
