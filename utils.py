import pandas as pd
import numpy as np
import streamlit as st
import io
from jinja2 import Environment, FileSystemLoader
from num2words import num2words
import os
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
from pypdf import PdfReader, PdfWriter
import tempfile
import shutil
import subprocess

# Initialize Jinja2 environment
env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")), cache_size=0)
env.filters['strptime'] = lambda s, fmt: datetime.strptime(s, fmt) if s else None

# Configure wkhtmltopdf with better error handling
wkhtmltopdf_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"

# Check if wkhtmltopdf exists
if not os.path.exists(wkhtmltopdf_path):
    raise FileNotFoundError(f"wkhtmltopdf not found at: {wkhtmltopdf_path}")

config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)

def number_to_words(number):
    try:
        return num2words(int(number), lang="en_IN").title()
    except:
        return str(number)

def make_gst_even(value):
    """
    Round GST value to nearest even number (no decimals).
    
    Args:
        value: The value to round
    
    Returns:
        Nearest even number as integer
    """
    try:
        # Round to nearest integer first
        rounded = round(value)
        
        # If rounded value is odd, adjust to nearest even
        if rounded % 2 != 0:
            return rounded + 1 if rounded % 2 != 0 else rounded
        return rounded
    except Exception as e:
        print(f"Error in make_gst_even: {str(e)}")
        return int(value)

def calculate_deductions(payable_amount, bill_type, is_first_bill=False):
    """
    Calculate all deductions based on bill type and amount.
    
    Args:
        payable_amount: Total payable amount before deductions
        bill_type: "Running Bill" or "Final Bill"
        is_first_bill: Boolean indicating if this is the first bill
    
    Returns:
        Dictionary containing all deduction amounts as integers
    """
    try:
        deductions = {
            'sd_amount': 0,
            'it_amount': 0,
            'gst_amount': 0,
            'lc_amount': 0,
            'recovery_deposit_v': 0,
            'total_deductions': 0,
            'by_cheque': 0,
            'cheque_amount_words': ''
        }
        
        # Calculate deductions
        if bill_type == "Running Bill":
            # For running bills, only calculate SD and IT
            deductions['sd_amount'] = make_gst_even(0.1 * payable_amount)
            deductions['it_amount'] = make_gst_even(0.02 * payable_amount)
            
            # For first bill, no SD deduction
            if is_first_bill:
                deductions['sd_amount'] = 0
        
        elif bill_type == "Final Bill":
            # For final bill, calculate all deductions
            deductions['sd_amount'] = make_gst_even(0.1 * payable_amount)
            deductions['it_amount'] = make_gst_even(0.02 * payable_amount)
            deductions['gst_amount'] = make_gst_even(0.02 * payable_amount)
            deductions['lc_amount'] = make_gst_even(0.01 * payable_amount)
        
        # Calculate total deductions
        deductions['total_deductions'] = sum([
            deductions['sd_amount'],
            deductions['it_amount'],
            deductions['gst_amount'],
            deductions['lc_amount'],
            deductions['recovery_deposit_v']
        ])
        
        # Calculate payment by cheque
        deductions['by_cheque'] = max(0, int(payable_amount) - deductions['total_deductions'])
        deductions['cheque_amount_words'] = number_to_words(deductions['by_cheque'])
        
        return deductions
    except Exception as e:
        print(f"Error in calculate_deductions: {str(e)}")
        return {
            'sd_amount': 0,
            'it_amount': 0,
            'gst_amount': 0,
            'lc_amount': 0,
            'recovery_deposit_v': 0,
            'total_deductions': 0,
            'by_cheque': int(payable_amount),
            'cheque_amount_words': number_to_words(int(payable_amount))
        }

def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type, amount_paid_last_bill, is_first_bill, user_inputs):
    """
    Process bill data from Excel sheets and prepare data for templates.
    
    Args:
        ws_wo: Work Order sheet DataFrame
        ws_bq: Bill Quantity sheet DataFrame
        ws_extra: Extra Items sheet DataFrame
        premium_percent: Premium percentage
        premium_type: Premium type ("Fixed" or "Percentage")
        amount_paid_last_bill: Amount paid in previous bill
        is_first_bill: Boolean indicating if this is the first bill
        user_inputs: Dictionary of user inputs from the form
    
    Returns:
        Tuple of data dictionaries for each document
    """
    try:
        # Initialize data structures
        first_page_data = {
            "header": [],
            "items": [],
            "totals": {
                "work_order_total": 0,
                "premium": {
                    "percent": premium_percent,
                    "amount": 0
                },
                "bill_amount": 0,
                "grand_total": 0
            }
        }
        
        last_page_data = {
            "payable_amount": 0,
            "amount_words": "",
            "amount_paid_last_bill": 0 if is_first_bill else int(amount_paid_last_bill),
            "payment_now": 0,
            "certificate_items": [],
            "bill_type": user_inputs.get("bill_type", "Running Bill"),
            "bill_number": user_inputs.get("bill_number", "First"),
            "last_bill": user_inputs.get("last_bill", "Not Applicable")
        }
        
        deviation_data = {
            "items": [],
            "summary": {
                "work_order_total": 0,
                "bill_amount": 0,
                "premium_percent": 0,
                "premium_type": "",
                "premium_amount": 0,
                "total_deviation": 0,
                "deviation_percentage": 0,
                "overall_excess": 0
            }
        } if user_inputs.get("bill_type", "Running Bill") == "Final Bill" else None
        
        extra_items_data = {
            "header": [],
            "items": [],
            "totals": {
                "work_order_total": 0,
                "extra_items_total": 0,
                "grand_total": 0
            }
        }
        
        note_sheet_data = {
            "notes": [],
            "deductions": {},
            "header": {},
            "work_order_amount": 0,
            "totals": {
                "work_order_total": 0,
                "bill_amount": 0,
                "grand_total": 0
            }
        }
        
        # Process header data
        header_data = ws_wo.iloc[:19, :7].replace(np.nan, "").values.tolist()
        for row in header_data:
            for i, val in enumerate(row):
                if isinstance(val, (pd.Timestamp, datetime)):
                    row[i] = val.strftime("%d-%m-%Y")
        first_page_data["header"] = header_data
        extra_items_data["header"] = header_data[:16]
        
        # Add user inputs to header
        note_sheet_data["header"] = {
            "agreement_no": user_inputs.get("agreement_no", ""),
            "name_of_work": user_inputs.get("work_name", ""),
            "name_of_firm": user_inputs.get("contractor_name", ""),
            "date_commencement": user_inputs.get("start_date", ""),
            "date_completion": user_inputs.get("completion_date", ""),
            "actual_completion": user_inputs.get("actual_completion_date", ""),
            "order_date": user_inputs.get("order_date", ""),
            "bill_type": user_inputs.get("bill_type", ""),
            "bill_number": user_inputs.get("bill_number", ""),
            "last_bill": user_inputs.get("last_bill", "")
        }
        
        # Process work order items
        work_order_total = 0
        last_row_wo = ws_wo.shape[0]
        for i in range(21, last_row_wo):
            qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else 0
            rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else 0
            
            try:
                qty = float(qty_raw) if isinstance(qty_raw, (int, float)) else \
                    float(qty_raw.strip().replace(',', '').replace(' ', '')) if isinstance(qty_raw, str) and qty_raw.strip() else 0
                
                rate = float(rate_raw) if isinstance(rate_raw, (int, float)) else \
                    float(rate_raw.strip().replace(',', '').replace(' ', '')) if isinstance(rate_raw, str) and rate_raw.strip() else 0
                
                amount = int(round(qty * rate))
                work_order_total += amount
                
                item = {
                    "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                    "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                    "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
                    "quantity": qty,
                    "rate": int(rate),
                    "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
                    "amount": amount,
                    "is_divider": False
                }
                first_page_data["items"].append(item)
            except Exception as e:
                print(f"Error processing work order item: {str(e)}")
                continue
        
        note_sheet_data["work_order_amount"] = work_order_total
        first_page_data["totals"]["work_order_total"] = work_order_total
        extra_items_data["totals"]["work_order_total"] = work_order_total
        note_sheet_data["totals"]["work_order_total"] = work_order_total
        
        # Process extra items
        try:
            first_page_data["items"].append({"description": "Extra Items", "is_divider": True})
            last_row_extra = ws_extra.shape[0]
            
            if last_row_extra > 6:  # Check if there are any extra items
                extra_items_total = 0
                for j in range(6, last_row_extra):
                    if ws_extra.shape[1] <= 5:
                        st.error(f"Extra Items sheet has insufficient columns: {ws_extra.shape[1]}")
                        break
                    
                    try:
                        serial_no = str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else ""
                        remark = str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else ""
                        description = str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else ""
                        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
                        unit = str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else ""
                        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None
                        
                        if qty_raw is not None and rate_raw is not None:
                            qty = float(qty_raw) if isinstance(qty_raw, (int, float)) else \
                                float(qty_raw.strip().replace(',', '').replace(' ', '')) if isinstance(qty_raw, str) and qty_raw.strip() else 0
                            
                            rate = float(rate_raw) if isinstance(rate_raw, (int, float)) else \
                                float(rate_raw.strip().replace(',', '').replace(' ', '')) if isinstance(rate_raw, str) and rate_raw.strip() else 0
                            
                            amount = int(round(qty * rate))
                            extra_items_total += amount
                            
                            item = {
                                "serial_no": serial_no,
                                "description": description,
                                "unit": unit,
                                "quantity": qty,
                                "rate": int(rate),
                                "remark": remark,
                                "amount": amount,
                                "is_divider": False
                            }
                            first_page_data["items"].append(item)
                            extra_items_data["items"].append(item)
                    except Exception as e:
                        print(f"Error processing extra item: {str(e)}")
                        continue
                
                extra_items_data["totals"]["extra_items_total"] = extra_items_total
                extra_items_data["totals"]["grand_total"] = work_order_total + extra_items_total
            else:
                # No extra items found
                first_page_data["items"].append({"description": "No Extra Items", "amount": 0, "is_divider": False})
                extra_items_data["items"].append({"description": "No Extra Items", "amount": 0, "is_divider": False})
                extra_items_data["totals"]["extra_items_total"] = 0
                extra_items_data["totals"]["grand_total"] = work_order_total
        except Exception as e:
            print(f"Error processing extra items section: {str(e)}")
            first_page_data["items"].append({"description": "No Extra Items", "amount": 0, "is_divider": False})
            extra_items_data["items"].append({"description": "No Extra Items", "amount": 0, "is_divider": False})
            extra_items_data["totals"]["extra_items_total"] = 0
            extra_items_data["totals"]["grand_total"] = work_order_total
        
        # Calculate totals
        first_page_data["totals"]["bill_amount"] = work_order_total
        
        # Initialize premium structure
        first_page_data["totals"]["premium"] = {
            "percent": premium_percent,
            "amount": 0
        }
        
        # Calculate premium if applicable
        if premium_percent > 0:
            if premium_type == "Fixed":
                first_page_data["totals"]["premium"]["amount"] = int(premium_percent)
            else:
                first_page_data["totals"]["premium"]["amount"] = int(round((premium_percent / 100) * work_order_total))
            first_page_data["totals"]["bill_amount"] = int(first_page_data["totals"]["bill_amount"] + first_page_data["totals"]["premium"]["amount"])
        
        first_page_data["totals"]["grand_total"] = first_page_data["totals"]["bill_amount"]
        
        # Calculate deductions
        last_page_data["payable_amount"] = first_page_data["totals"]["bill_amount"]
        last_page_data["amount_words"] = number_to_words(last_page_data["payable_amount"])
        
        # Get deductions based on bill type
        deductions = calculate_deductions(
            last_page_data["payable_amount"],
            user_inputs.get("bill_type", "Running Bill"),
            is_first_bill
        )
        
        # Update last page data with deductions
        last_page_data.update(deductions)
        
        # Add certificate items
        last_page_data["certificate_items"] = [
            {"name": "S.D.", "percentage": 10, "value": last_page_data["sd_amount"]},
            {"name": "I.T.", "percentage": 2, "value": last_page_data["it_amount"]},
            {"name": "GST", "percentage": 2, "value": last_page_data["gst_amount"]},
            {"name": "L.C.", "percentage": 1, "value": last_page_data["lc_amount"]}
        ]
        
        # Only process deviation data for final bills
        if user_inputs.get("bill_type", "Running Bill") == "Final Bill":
            # Calculate deviation data
            deviation_data["summary"]["work_order_total"] = work_order_total
            deviation_data["summary"]["bill_amount"] = first_page_data["totals"]["bill_amount"]
            deviation_data["summary"]["premium_percent"] = premium_percent
            deviation_data["summary"]["premium_type"] = premium_type
            deviation_data["summary"]["premium_amount"] = first_page_data["totals"]["premium"]["amount"]
            
            # Calculate overall deviation
            deviation_data["summary"]["total_deviation"] = first_page_data["totals"]["bill_amount"] - work_order_total
            deviation_data["summary"]["deviation_percentage"] = (deviation_data["summary"]["total_deviation"] / work_order_total) * 100 if work_order_total > 0 else 0
            deviation_data["summary"]["overall_excess"] = deviation_data["summary"]["total_deviation"]
        
        return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data
        
    except Exception as e:
        st.error(f"Error processing bill: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")
        return None, None, None, None, None

def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount, note_sheet_data):
    """
    Generate note sheet content.
    
    Args:
        payable_amount: Total payable amount
        work_order_amount: Work order amount
        extra_item_amount: Extra items amount
        note_sheet_data: Dictionary containing note sheet data
    """
    try:
        # Calculate deductions
        sd_amount = make_gst_even(0.1 * payable_amount)
        it_amount = make_gst_even(0.02 * payable_amount)
        gst_amount = make_gst_even(0.02 * payable_amount)
        lc_amount = make_gst_even(0.01 * payable_amount)
        total_deductions = sd_amount + it_amount + gst_amount + lc_amount
        by_cheque = payable_amount - total_deductions
        
        # Format the data for the template
        data = {
            'payable_amount': f"{payable_amount:,.2f}",
            'total_123': f"{payable_amount:,.2f}",
            'balance_4_minus_5': f"{payable_amount:,.2f}",
            'amount_paid_last_bill': f"{note_sheet_data.get('amount_paid_last_bill', 0):,.2f}",
            'payment_now': f"{by_cheque:,.2f}",
            'by_cheque': f"{by_cheque:,.2f}",
            'cheque_amount_words': number_to_words(by_cheque),
            'total_recovery': f"{total_deductions:,.2f}",
            'certificate_items': [
                {'name': 'S.D.', 'percentage': 10, 'value': f"{sd_amount:,.2f}"},
                {'name': 'I.T.', 'percentage': 2, 'value': f"{it_amount:,.2f}"},
                {'name': 'GST', 'percentage': 2, 'value': f"{gst_amount:,.2f}"},
                {'name': 'L.C.', 'percentage': 1, 'value': f"{lc_amount:,.2f}"}
            ]
        }
        
        return data
        
    except Exception as e:
        st.error(f"Error generating note sheet: {str(e)}")
        return None

def generate_pdf(html_content, output_path=None):
    """
    Generate PDF from HTML content
    
    Args:
        html_content: HTML content as string
        output_path: Optional output path for the PDF
    """
    try:
        # Create temporary directory for PDF generation
        with tempfile.TemporaryDirectory() as temp_dir:
            # Create temporary HTML file
            temp_html = os.path.join(temp_dir, "temp.html")
            with open(temp_html, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Try to find wkhtmltopdf in common locations
            possible_paths = [
                r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe",
                r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe",
                r"C:\Program Files\wkhtmltopdf\wkhtmltopdf.exe",
                r"C:\Program Files (x86)\wkhtmltopdf\wkhtmltopdf.exe"
            ]
            
            wkhtmltopdf_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    wkhtmltopdf_path = path
                    break
            
            if not wkhtmltopdf_path:
                raise FileNotFoundError("wkhtmltopdf executable not found. Please install it from: https://wkhtmltopdf.org/downloads.html")
            
            # Configure wkhtmltopdf
            config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
            
            # Generate PDF
            pdf_bytes = pdfkit.from_file(
                temp_html,
                False,  # Output to BytesIO
                configuration=config,
                options={
                    'page-size': 'A4',
                    'margin-top': '0.25in',
                    'margin-bottom': '0.25in',
                    'margin-left': '0.25in',
                    'margin-right': '0.5in',
                    'encoding': "UTF-8",
                    'quiet': "",
                    'no-outline': None,
                    'enable-local-file-access': None,
                    'disable-smart-shrinking': None,
                    'dpi': 300,
                    'javascript-delay': "1000",
                    'no-stop-slow-scripts': None,
                    'load-error-handling': "ignore"
                }
            )
            
            # If output_path is provided, write to file
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(pdf_bytes)
                return None
            
            # Return BytesIO object
            return io.BytesIO(pdf_bytes)
            
    except Exception as e:
        error_msg = f"Error generating PDF: {str(e)}"
        print(f"Error details: {traceback.format_exc()}")
        raise ValueError(error_msg) from e

def combine_pdfs(pdf_paths, output_path):
    """
    Combine multiple PDFs into one.
    
    Args:
        pdf_paths: List of paths to PDF files
        output_path: Path where combined PDF should be saved
    """
    try:
        # Validate output path
        if not output_path:
            raise ValueError("Output path cannot be empty")
        
        # Validate PDF paths
        if not pdf_paths:
            raise ValueError("No PDFs to combine")
        
        # Create temporary directory for PDFs
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Combine PDFs using pdftk
            cmd = [
                "pdftk",
                *pdf_paths,
                "cat",
                "output",
                output_path
            ]
            
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            except subprocess.CalledProcessError as e:
                raise RuntimeError(f"Error combining PDFs: {e.stderr}\nCommand: {' '.join(cmd)}")
            except FileNotFoundError:
                raise RuntimeError("pdftk executable not found. Please install it from https://www.pdflabs.com/tools/pdftk-the-pdf-toolkit/")
            
            # Verify PDF was created
            if not os.path.exists(output_path):
                raise IOError(f"Failed to create combined PDF: {output_path}")
                
            # Verify PDF size (should be at least 1KB)
            if os.path.getsize(output_path) < 1024:
                raise IOError(f"Generated combined PDF is too small: {output_path}")
                
            return True
            
        finally:
            # Clean up temporary files
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
                
    except Exception as e:
        print(f"Error in combine_pdfs: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

def create_word_doc(sheet_name, data, doc_path, header_data=None):
    """
    Create Word document from template.
    
    Args:
        sheet_name: Name of the sheet/template
        data: Dictionary containing data for template
        doc_path: Path to save the Word document
        header_data: Optional header data
    """
    try:
        # Create new Word document
        doc = Document()
        
        # Add header if provided
        if header_data:
            header = doc.add_heading("Contractor Bill", 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add header data as table
            header_table = doc.add_table(rows=1, cols=2)
            header_table.style = 'Table Grid'
            
            for row in header_data:
                row_cells = header_table.add_row().cells
                row_cells[0].text = str(row[0])
                row_cells[1].text = str(row[1])
        
        # Add main content based on sheet name
        if sheet_name == "first_page":
            # Add bill items table
            items_table = doc.add_table(rows=1, cols=7)
            items_table.style = 'Table Grid'
            
            # Add headers
            headers = ['S.No', 'Description', 'Unit', 'Quantity', 'Rate', 'Amount', 'Remark']
            for i, header in enumerate(headers):
                cell = items_table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Add items
            for item in data['items']:
                if item['is_divider']:
                    doc.add_paragraph("-" * 80)
                else:
                    row = items_table.add_row().cells
                    row[0].text = str(item['serial_no'])
                    row[1].text = str(item['description'])
                    row[2].text = str(item['unit'])
                    row[3].text = f"{item['quantity']:.2f}"
                    row[4].text = f"{item['rate']:.2f}"
                    row[5].text = f"{item['amount']:.2f}"
                    row[6].text = str(item['remark'])
        
        # Save document
        doc.save(doc_path)
        
    except Exception as e:
        st.error(f"Error creating Word document for {sheet_name}: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")